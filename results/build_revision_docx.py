"""Build the paper revision + Methods/Data-Sources supplement as a .docx.

This does NOT overwrite the authors' manuscript. It produces an editorial pass:
  Part A — the framing fix (isolate germline editing) + section-by-section proposed edits
  Part B — a complete, drop-in "Methods and Data Sources" section
  Part C — an updated Table 1 and a change log
All quantitative claims are pulled from results/paper_numbers.json so the document
stays internally consistent with the computed pipeline.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
PN = json.loads((ROOT / "paper_numbers.json").read_text())
_LIB = json.loads((ROOT.parent / "app" / "public" / "data" / "library.json").read_text())
TIERS = _LIB["rollup"]["tiers"]
N_CATALOGUE = TIERS["all"]["n_diseases"]
N_CORE = TIERS["core"]["n_diseases"]
N_RARE = TIERS["rare"]["n_diseases"]
CITED_SHARE_ALL = TIERS["all"]["cited_incidence_share_by_count"]


def med(key):
    return PN[key]["median"]


def ci(key):
    return PN[key]["ci95"]


def M(x):
    return f"{x/1e6:.2f}M"


def k(x):
    return f"{x/1e3:.0f},000" if x >= 1000 else f"{x:.0f}"


def pct(x, d=1):
    return f"{x*100:.{d}f}%"


ACCENT = RGBColor(0x1F, 0x49, 0x7D)
GREY = RGBColor(0x55, 0x55, 0x55)
VIOLET = RGBColor(0x5B, 0x21, 0xB6)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def h3(t):
    return doc.add_heading(t, level=3)


def para(t, italic=False, color=None, size=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def rich(parts):
    """parts: list of (text, {'b':bool,'i':bool,'color':RGBColor})."""
    p = doc.add_paragraph()
    for text, fmt in parts:
        r = p.add_run(text)
        r.bold = fmt.get("b", False)
        r.italic = fmt.get("i", False)
        if fmt.get("color"):
            r.font.color.rgb = fmt["color"]
    return p


def bullet(t, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.add_run(t)
    return p


def quote(label, t):
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = GREY
    r2 = p.add_run(t)
    r2.italic = True
    r2.font.color.rgb = GREY
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def before(t):
    quote("BEFORE —", t)


def after(t):
    p = doc.add_paragraph()
    r = p.add_run("AFTER — ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(t)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


# ---------------------------------------------------------------- title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Genetic Medicine — Revision Pass and Methods/Data-Sources Supplement")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run(
    "Proposed editorial revisions to the draft manuscript, with all quantitative claims "
    "regenerated from the open analysis pipeline. Not a replacement manuscript — a marked pass "
    "intended to be reviewed and merged by the authors."
)
rs.italic = True
rs.font.color.rgb = GREY
rs.font.size = Pt(10)

para("")
rich([
    ("Two things drive this pass. ", {"b": True}),
    ("First, the paper's central move — that existing genetic medicine already reaches the "
     "overwhelming majority of serious genetic disease, and germline editing is uniquely needed "
     "only for a tiny residual — is blunted wherever germline editing is allowed to blur into "
     "“what genetic medicine can already do.” The edits below keep germline editing "
     "categorically separate everywhere: it is not one of the four existing tools, and it is not a "
     "post-birth treatment modality (not even a cousin of somatic gene therapy). ", {}),
    ("Second, ", {"b": True}),
    ("because part of the contribution is a new framing of genetic-medicine impact, the analysis "
     "decisions — which data we used, which we deliberately did not, and why — need to be stated "
     "precisely. Part B is a drop-in Methods and Data Sources section that does this.", {}),
])

# ---------------------------------------------------------------- Part A
h1("Part A — The framing fix: isolate germline editing")

h2("A.0  The one distinction the paper must never blur")
para(
    "There are three categorically different things in play, and the argument only lands if the "
    "reader can always tell which is which:")
rich([("1. The existing reproductive/diagnostic tools — ", {"b": True}),
      ("preconception carrier screening (CS), IVF with preimplantation genetic testing (PGT), "
       "prenatal diagnosis (PND), and newborn screening (NBS). These prevent or pre-empt an "
       "affected birth. They exist today and are validated at national scale.", {})])
rich([("2. Existing post-birth therapies — ", {"b": True}),
      ("surgery, drugs, dietary/cofactor management, enzyme replacement, transplant, and "
       "somatic gene/cell therapy. These treat a child who is already born. They also exist "
       "today. Somatic gene therapy edits or supplements genes in one patient's body tissue; it "
       "is not heritable and it is not germline editing.", {})])
rich([("3. Germline (embryo/germ-cell) editing — ", {"b": True, "color": VIOLET}),
      ("a heritable change made before birth. This is the future, contested technology the paper "
       "is really about. It belongs to neither (1) nor (2). It is tracked separately, as a "
       "residual, and never counted inside “existing tools” or “treatable.”", {})])
para(
    "Recommendation: adopt these three labels verbatim and use them consistently. Where the draft "
    "says “existing methods,” specify whether it means (1), (2), or (1)+(2) — it should "
    "almost always mean (1)+(2), i.e. everything that is not germline editing.",
    italic=True)

h2("A.1  Abstract")
before(
    "…over 99% of serious heritable disease in presently understood cause is in principle "
    "preventable or substantially mitigable using existing tools… By contrast, germline "
    "embryo editing is relevant only to a narrow residual category…")
after(
    "Using global birth and disease-burden estimates, and a purpose-built library of "
    f"{N_CATALOGUE} serious genetic conditions mapped to their causal genes and "
    "to each intervention, we show that the existing tools of genetic medicine — carrier "
    "screening, IVF with preimplantation genetic testing, prenatal diagnosis, and newborn "
    "screening, together with today's post-birth therapies — already address on the order of "
    f"{pct(med('addressable_share_of_serious__permissive'),0)} of serious genetic disease "
    f"(95% credible interval {pct(ci('addressable_share_of_serious__permissive')[0],0)}–"
    f"{pct(ci('addressable_share_of_serious__permissive')[1],0)}). Germline editing is a "
    "categorically distinct intervention — neither one of these tools nor a post-birth treatment — "
    "and is uniquely required only for a narrow residual: couples for whom no unaffected embryo can "
    f"be selected, on the order of {k(med('uniquely_editable_births_per_year__permissive'))} births "
    f"per year ({pct(med('uniquely_editable_share_of_serious__permissive'),1)} of serious genetic "
    "disease) even under deliberately permissive assumptions. We keep germline editing separate "
    "from existing tools throughout, so that its true — small but real — domain is neither inflated "
    "by borrowing the reach of screening nor dismissed.")

h2("A.2  Section II (Taxonomy) — add a one-line firewall")
para(
    "The taxonomy already lists somatic vs germ-line editing correctly. Add one sentence under "
    "“Therapeutic” so a reader cannot later conflate somatic gene therapy with germline "
    "editing when the impact numbers arrive:")
after(
    "A note on vocabulary: somatic genetic medicines (siRNA/mRNA/ASO, ex-vivo and in-vivo somatic "
    "DNA editing) act on the cells of a person already born and are not heritable. Germline "
    "editing acts on the embryo or germ cell and is heritable. Throughout the impact analysis we "
    "count somatic therapies among existing post-birth treatments and germline editing entirely "
    "separately; the two must not be summed.")

h2("A.3  Section IV (Analyzing Impact) — the core rewrite")
para("This is where the blurring does the most damage. Three targeted changes.")

h3("(a) Replace the paragraph that sets up the quantification")
before(
    "Now we can quantify how much of the global genetic disease burden these four methods could "
    "prevent, and then calculate the remainder for which embryo editing would be the only solution")
after(
    "We can now quantify how much of the serious genetic-disease burden the existing tools address "
    "— splitting them by the type of action they take — and then isolate the residual for which "
    "germline editing, and only germline editing, would add something. We keep three ledgers "
    "separate: (i) affected births prevented before birth by carrier screening, PGT, or prenatal "
    "diagnosis; (ii) affected children whose disease is meaningfully modified after birth by an "
    "existing therapy (newborn-screening-linked treatment, drugs, enzyme replacement, somatic "
    "gene/cell therapy); and (iii) the germline-editing-only residual. A case is assigned to (iii) "
    "only if it is reachable by neither (i) nor (ii).")

h3("(b) Numbers paragraph — update to the regenerated figures")
before(
    "…embryo editing's theoretical domain amounts to about 1.8% of serious genetic disease "
    "and 0.1% of all births. Everything else – about 98.2% – lies within reach of the "
    "four existing methods…")
after(
    "Isolating the germline-editing-only residual makes the picture stark. Existing tools address "
    f"{pct(med('addressable_share_of_serious__permissive'),1)} of serious genetic disease "
    f"(95% CrI {pct(ci('addressable_share_of_serious__permissive')[0],1)}–"
    f"{pct(ci('addressable_share_of_serious__permissive')[1],1)}) even under permissive assumptions "
    "about where editing might help. The residual uniquely needing germline editing is "
    f"{pct(med('uniquely_editable_share_of_serious__permissive'),1)} of serious genetic disease — "
    f"about {k(med('uniquely_editable_births_per_year__permissive'))} births per year, roughly "
    f"{pct(med('uniquely_editable_share_of_serious__permissive')*med('serious_share_of_all_births'),3)} "
    "of all births. Under a strict definition (only couples with literally no selectable unaffected "
    f"embryo) the residual falls to about {k(med('uniquely_editable_births_per_year__strict'))} "
    f"births per year ({pct(med('uniquely_editable_share_of_serious__strict'),1)} of serious "
    "disease). The no-selectable-embryo core of that residual — the least disputable case for "
    f"editing — is on the order of {k(med('s1_no_selectable_embryo_births_per_year'))} births per "
    "year. The direction of the argument does not depend on which bound one takes: on any "
    "assumption, existing tools reach the overwhelming majority and germline editing reaches a "
    "sliver.")

h3("(c) Add the somatic-vs-germline guardrail to the resistance section")
para(
    "In the cardiovascular passage, the draft's own logic already shows somatic PCSK9 inhibitors "
    "matching the benefit — reinforce that this is an existing (2)-type therapy, not a case for (3):")
after(
    "Note the structure of this example: the benefit is already delivered by a somatic therapy "
    "(statins, somatic PCSK9 inhibitors). That is existing genetic/pharmacologic medicine acting "
    "after birth — not germline editing. It is precisely because an existing post-birth therapy "
    "matches the effect that the germline-resistance case here is weak.")

# ---------------------------------------------------------------- Table 1
h2("A.4  Table 1 — rebuilt so existing tools and germline editing are different columns")
para(
    "The draft references a prevention table but leaves it as a caption. Below is a drop-in table, "
    "populated from the regenerated numbers, that structurally separates existing tools from "
    "germline editing rather than listing editing as one more method.")

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, txt in enumerate([
        "Serious genetic disease (per year)",
        "Addressed by existing tools (CS · PGT · PND · NBS + post-birth therapy)",
        "Uniquely needs germline editing",
        "Share"]):
    hdr[i].paragraphs[0].add_run(txt).bold = True

serious = med("serious_genetic_births_per_year")
mono = med("monogenic_births_per_year")
multi = med("multifactorial_births_per_year")
edit_perm = med("uniquely_editable_births_per_year__permissive")
edit_strict = med("uniquely_editable_births_per_year__strict")
s1 = med("s1_no_selectable_embryo_births_per_year")

rows = [
    ("All serious genetic disease", f"~{M(serious - edit_perm)} ({pct(med('addressable_share_of_serious__permissive'),1)})",
     f"~{k(edit_perm)} (permissive)", pct(med('uniquely_editable_share_of_serious__permissive'),1)),
    ("  of which monogenic", f"~{M(mono)} total; near-all screenable/selectable", f"no-selectable-embryo core ~{k(s1)}", pct(s1/serious,2)),
    ("  of which multifactorial", f"~{M(multi)} total; risk-reducible, not single-edit", "≈0 today (needs multi-locus editing)", "—"),
    ("Strict residual (no selectable embryo only)", f"~{M(serious - edit_strict)}", f"~{k(edit_strict)}", pct(med('uniquely_editable_share_of_serious__strict'),1)),
]
for a, b, c, d in rows:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(a)
    cells[1].paragraphs[0].add_run(b)
    rc = cells[2].paragraphs[0].add_run(c)
    rc.font.color.rgb = VIOLET
    cells[3].paragraphs[0].add_run(d)

para("")
para(
    "All values are Monte-Carlo medians over 20,000 draws; see Part B for the credible intervals "
    "and the severity/attribution assumptions. "
    "The middle column is everything that is NOT germline editing; the third column is germline "
    "editing alone. They are never added together.",
    italic=True, color=GREY, size=9)

# ---------------------------------------------------------------- Part B
doc.add_page_break()
h1("Part B — Methods and Data Sources (drop-in section)")
para(
    "This section is written to be inserted into the manuscript (as a Methods section or a "
    "supplementary appendix). It states the framing, every data source used, the sources "
    "deliberately not used and why, and the modelling and accounting choices. A reproducible "
    "pipeline and an interactive version of every figure accompany the paper.",
    italic=True, color=GREY)

h2("B.1  What is new in the framing")
para(
    "We treat genetic-medicine impact as a library problem rather than a single headline ratio. "
    "The unit of analysis is a serious genetic disease, mapped to (a) its causal gene(s) and mode "
    "of inheritance, (b) each existing intervention that applies to it (carrier screening, PGT, "
    "prenatal diagnosis, newborn screening) and the type of any existing post-birth therapy, and "
    "(c) whether germline editing would add anything the existing options cannot. Two design "
    "choices distinguish this from a burden estimate:")
bullet("Germline editing is never a member of the intervention set or the treatment taxonomy. It "
       "is computed as a residual — the cases reachable by no existing tool — so it can neither "
       "inherit the reach of screening nor be hidden inside “treatable.”")
bullet("Every disease carries a genetic-medicine status derived only from which interventions "
       "apply (preventable, treatable, preventable-and-treatable, detectable-only, or none), with "
       "no hidden weighting, so the headline split is auditable disease-by-disease.")

h2("B.2  Data sources used")
para("Each source is used for a specific quantity; none is used beyond what it measures.")

h3("Birth cohort")
bullet("United Nations, World Population Prospects 2024 — annual global births "
       f"(~{M(med('births_per_year'))}; 95% CrI {M(ci('births_per_year')[0])}–{M(ci('births_per_year')[1])}). "
       "Denominator for every per-birth figure.")

h3("Serious genetic-disease burden (top-down anchors)")
bullet("March of Dimes, Global Report on Birth Defects (2006) — ~6% of live births affected by "
       "serious congenital disorders of substantial genetic origin (~8 million/year), and the "
       "~94% share born in low- and middle-income countries.")
bullet("Modell & Darlison, Bull. WHO (2008) — global epidemiology of haemoglobin disorders and "
       "derived service indicators; anchors the monogenic serious-disease rates and the "
       "screening-programme reduction fractions.")
bullet("WHO, Congenital anomalies fact sheet (2023) — corroborating congenital-anomaly totals.")

h3("Condition-specific incidence (GBD 2023, IHME)")
bullet("Global Burden of Disease 2023 — birth/<1-year incidence for the highest-burden conditions "
       "used in the core catalogue: sickle cell disorders, congenital heart anomalies, neural tube "
       "defects, Down syndrome, and orofacial clefts. These replace textbook estimates with cited "
       "incidences for the conditions that dominate the totals.")

h3("Catalogue expansion and prevalence (Orphanet / Orphadata)")
bullet("Orphadata en_product9_prev (prevalence), en_product6 (gene associations), and "
       "en_product9_ages (inheritance). Used two ways, with different disciplines (see B.3).")

h3("Intervention effect sizes and costs")
bullet("ACMG (2021) carrier-screening practice resource; Harper et al. PGT outcomes; Wald (2018) "
       "and Nordic prenatal-programme reductions; Kirby & Browne (2021) newborn screening; WHO "
       "PMTCT (>98% prevention where implemented) and UNAIDS 2023 vertical-infection counts "
       f"(~{k(med('hiv_vertical_infections_per_year'))}/yr, residual after PMTCT "
       f"~{k(med('hiv_residual_after_pmtct'))}); published gene-therapy and screening-programme "
       "unit costs for the allocation comparison.")

h2("B.3  Sources deliberately NOT used (and why)")
para(
    "Being explicit about exclusions is part of the framing contribution: the residual attributed "
    "to germline editing is sensitive to exactly these choices.")
bullet("Single-country / regional Orphanet birth prevalences were NOT promoted into the core "
       "catalogue. Founder-effect and ascertainment differences make them poor global anchors "
       "(e.g. a national cystic-fibrosis figure, or sickle-cell on a single island). Orphanet was "
       "allowed to overwrite a core incidence only when it carried a Worldwide “prevalence at "
       "birth” value and the existing figure was a mere textbook estimate — 26 conditions. "
       "Single-country values appear only in the clearly flagged rare tier (B.4), with their "
       "geography attached.")
bullet("gnomAD allele-frequency auto-ingestion was deferred; allele frequencies for the "
       "no-selectable-embryo residual are taken from curated literature values instead of an "
       "automated pull, to avoid mis-mapping variants to disorders.")
bullet("Automated UN WPP and UNAIDS API pulls were unavailable in the build environment; the "
       "corresponding figures are taken from the published reports rather than a live endpoint.")
bullet("Enhancement, and most disease-resistance indications, are excluded from the prevention "
       "denominator. They are assessed separately (Section IV) and never folded into "
       "“genetic disease,” because doing so would let non-prevention aims borrow the "
       "moral weight of prevention — one of the paper's central cautions.")
bullet("Non-genetic congenital anomalies are excluded via the attribution stance (below), so the "
       "denominator is genetic disease, not all birth defects.")

h2("B.4  The disease library: a curated core plus an Orphanet rare tier")
rich([
    ("The catalogue has two tiers. A hand-curated ", {}),
    ("core", {"b": True}),
    (" of the highest-burden serious conditions carries the global numbers and is calibrated "
     "against the top-down burden model. An Orphanet-derived ", {}),
    ("rare tier", {"b": True}),
    (" adds every other Orphanet disorder with a cited birth prevalence, a known causal gene, and "
     "a clean monogenic inheritance mode — individually rare, collectively a long tail — so the "
     "catalogue answers disease-count questions without the rare tail distorting the "
     "burden-weighted headline (which is computed over the core alone). For the rare tier, "
     "intervention applicability is assigned by transparent rule, not case-by-case curation: "
     "carrier screening for recessive/X-linked conditions; PGT and prenatal diagnosis for any "
     "monogenic condition with a known gene; newborn treatment left uncredited pending curation "
     "(a conservative default that can only understate what genetic medicine offers). The core "
     f"holds {N_CORE} conditions and the rare tier {N_RARE}, for {N_CATALOGUE} in all; "
     f"{pct(CITED_SHARE_ALL,0)} of the full catalogue now rests on a cited incidence.", {}),
])

h2("B.5  The two independent estimates, and why they differ")
para(
    "We report a top-down parametric denominator and a bottom-up catalogue sum, deliberately kept "
    "separate. The top-down model samples birth counts, per-1,000 serious-disease rates, and an "
    "attribution stance from cited intervals to produce a calibrated total with credible "
    f"intervals (~{M(med('serious_genetic_births_per_year'))} serious births/yr; 95% CrI "
    f"{M(ci('serious_genetic_births_per_year')[0])}–{M(ci('serious_genetic_births_per_year')[1])}). "
    "The bottom-up catalogue sum is a lower bound that climbs toward the top-down total as the "
    "library grows. Presenting both makes the coverage gap explicit rather than hiding it behind a "
    "single number.")

h2("B.6  Modelling and accounting choices")
bullet("Severity threshold. Three definitions of “serious” (A: catastrophic/early-onset "
       "only; B: serious incl. later-onset and major metabolic/immune/cardiac; C: broad) — default "
       "B. Every headline is reported against the threshold used.")
bullet("Attribution stance. Inclusive (count a condition's full burden), heritability-weighted "
       "(scale by genetic contribution), or exclusive (monogenic/chromosomal only) — default "
       "inclusive. This is the single largest swing on the denominator and is exposed as a toggle.")
bullet("Monte-Carlo uncertainty. 20,000 draws; Beta priors for proportions, Lognormal for rates "
       "and costs; every ratio computed per-draw so credible intervals are correct rather than "
       "propagated by hand.")
bullet("Germline-editing residual. Two components: S1, couples with no selectable unaffected "
       "embryo (both parents affected by the same dominant condition, or a parent homozygous for a "
       "fully penetrant recessive), computed from allele frequency, penetrance, survival, "
       "assortative mating and consanguinity; and S2, cases where a single reliable edit would beat "
       "every existing option for a complex condition (≈0 today). Permissive vs strict bounds "
       "differ only in how generously S2 and partially-penetrant cases are credited to editing.")
bullet("Contested conditions. Congenital deafness is reported both in and out of the residual, "
       "since whether it counts as “serious disease” is genuinely contested; the headline "
       "uses the definition that excludes it, and its contribution is shown separately.")
bullet("Embryo accounting. Selection-based strategies discard affected embryos; germline editing "
       "does not. We track embryos created and affected embryos discarded per unaffected child "
       "under each strategy, since this is a substantive normative difference between selection and "
       "editing, not a modelling detail.")
bullet("Costing. Cost per affected birth prevented is compared for screening programmes "
       f"(~${k(med('cost_per_birth_prevented_screening'))}) versus an editing programme "
       f"(~${k(med('cost_per_birth_prevented_editing'))}) — roughly a "
       f"{med('cost_per_birth_prevented_editing')/med('cost_per_birth_prevented_screening'):.0f}× "
       "difference — with credible intervals in the supplement.")

h2("B.7  Reproducibility and limitations")
para(
    "The full pipeline (data ingestion, the disease library, the Monte-Carlo model, and every "
    "figure) is open and regenerates all numbers in this paper from cited inputs; an interactive "
    "version lets a reader move the severity threshold and attribution stance and watch every "
    "figure respond. Limitations: the bottom-up catalogue is a lower bound and still under-covers "
    "the rare monogenic tail; rare-tier interventions are rule-assigned, not individually "
    "vetted; single-country prevalences carry founder-effect bias; and the S2 (editing-superior) "
    "term for complex disease is model-dependent and, on current evidence, indistinguishable from "
    "zero.")

# ---------------------------------------------------------------- Part C
doc.add_page_break()
h1("Part C — Change log")
changes = [
    ("Abstract", "Rewritten to name the four existing tools + post-birth therapy as the "
     "“existing” set and place germline editing outside it; numbers updated."),
    ("Sec II Taxonomy", "One-sentence firewall added distinguishing somatic genetic medicine from "
     "germline editing."),
    ("Sec IV setup", "Quantification paragraph rewritten to keep three separate ledgers (prevent / "
     "treat / editing-only)."),
    ("Sec IV numbers", f"Updated to regenerated figures: existing tools "
     f"{pct(med('addressable_share_of_serious__permissive'),1)}, editing residual "
     f"{pct(med('uniquely_editable_share_of_serious__permissive'),1)} (permissive) down to "
     f"{pct(med('uniquely_editable_share_of_serious__strict'),1)} (strict)."),
    ("Sec IV resistance", "Cardiovascular passage tagged as a somatic (existing) therapy case, not "
     "a germline-editing case."),
    ("Table 1", "Rebuilt with existing tools and germline editing as separate columns."),
    ("Methods & Data Sources", "New drop-in section: framing, sources used, sources excluded and "
     "why, library tiers, dual estimates, modelling/accounting choices, reproducibility."),
    ("References", "Add: IHME GBD 2023; Orphanet/Orphadata (prevalence, genes, inheritance); "
     "UN WPP 2024. Existing Modell & Darlison, March of Dimes, ACMG, Wald, Kirby, UNAIDS, WHO "
     "retained."),
]
for a, b in changes:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(a + " — ")
    r.bold = True
    p.add_run(b)

out = ROOT / "Genetic_Medicine_Revision_and_Methods.docx"
doc.save(out)
print("wrote", out)
